from flask import Flask, render_template, request, send_file, redirect, url_for, flash, send_from_directory
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from docx.shared import Inches
from copy import deepcopy
from pathlib import Path
from datetime import datetime
import json
import os
import re
import shutil
import uuid

try:
    from PIL import Image
    from pillow_heif import register_heif_opener

    register_heif_opener()
    HEIC_SUPPORTED = True
except Exception:
    HEIC_SUPPORTED = False


app = Flask(__name__)
app.secret_key = "atvsld-local-secret"

BASE_DIR = Path(__file__).resolve().parent

login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = "login"

USERS_FILE = BASE_DIR / "users.json"


class User(UserMixin):
    def __init__(self, username, password_hash, can_edit=False):
        self.id = username
        self.username = username
        self.password_hash = password_hash
        self.can_edit = can_edit


def load_users():
    if not USERS_FILE.exists():
        return {}

    try:
        with open(USERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_users(users):
    with open(USERS_FILE, "w", encoding="utf-8") as f:
        json.dump(users, f, ensure_ascii=False, indent=2)


def get_user(username):
    users = load_users()
    data = users.get(username)

    if not data:
        return None

    return User(
        username=username,
        password_hash=data.get("password_hash", ""),
        can_edit=data.get("can_edit", False)
    )


@login_manager.user_loader
def load_user(username):
    return get_user(username)


def ensure_admin():
    users = load_users()

    if "admin" not in users:
        users["admin"] = {
            "password_hash": generate_password_hash("123456"),
            "can_edit": True,
            "is_admin": True
        }

        save_users(users)


BASE_DIR = Path(__file__).resolve().parent

USERS_FILE = BASE_DIR / "users.json"

TEMPLATE = BASE_DIR / "Bien_ban_ATVSLD_thang07_2026_VHIALY fn.docx"

DATA_DIR = BASE_DIR / "data"
REPORTS_DIR = BASE_DIR / "reports"
OUTPUT_DIR = BASE_DIR / "output"

DATA_DIR.mkdir(exist_ok=True)
REPORTS_DIR.mkdir(exist_ok=True)
OUTPUT_DIR.mkdir(exist_ok=True)


# Tạo tài khoản admin mặc định
ensure_admin()

def edit_required():
    if not current_user.is_authenticated:
        flash("Anh/chị cần đăng nhập để thực hiện chức năng này.", "error")
        return redirect(url_for("login"))

    if not current_user.can_edit:
        flash("Tài khoản này không có quyền chỉnh sửa.", "error")
        return redirect(url_for("history"))

    return None


@app.route("/login", methods=["GET", "POST"])
def login():
    if current_user.is_authenticated:
        return redirect(url_for("index"))

    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")

        user = get_user(username)

        if user and check_password_hash(user.password_hash, password):
            login_user(user)
            return redirect(url_for("index"))

        flash("Tên đăng nhập hoặc mật khẩu không đúng.", "error")

    return render_template("login.html")


@app.route("/logout")
@login_required
def logout():
    logout_user()
    return redirect(url_for("login"))


ALLOWED_EXT = {
    "jpg",
    "jpeg",
    "png",
    "webp",
    "gif",
    "heic",
    "heif"
}

DEFAULT_GROUPS = [
    {"stt":"1","title":"Việc thực hiện các quy định về ATVSLĐ; khai báo, điều tra, thống kê tai nạn lao động; đánh giá nguy cơ rủi ro về ATVSLĐ; huấn luyện về ATVSLĐ, xây dựng và thực hiện VHAT","rows":[
        ["1.1","Việc thực hiện các quy định về ATVSLĐ; khai báo, điều tra, thống kê tai nạn lao động","Thực hiện đầy đủ các quy định về ATVSLĐ; không xảy ra tai nạn lao động, không có vụ việc phải khai báo hoặc điều tra","Không phát sinh kiến nghị sau kiểm tra",""] ,
        ["1.2","Đánh giá nguy cơ rủi ro về ATVSLĐ","Đã thực hiện đánh giá nguy cơ rủi ro về ATVSLĐ theo quy định","Không phát sinh kiến nghị sau kiểm tra",""] ,
        ["1.3","Huấn luyện về ATVSLĐ, xây dựng và thực hiện VHAT","- Người lao động được huấn luyện đầy đủ về ATVSLĐ và chấp hành tốt các quy định về Văn hóa an toàn.\n- Phân xưởng đã Diễn tập XLSC; ứng cứu khẩn cấp ATVSLĐ; chữa cháy và cứu nạn, cứu hộ tại các Nhà máy Ialy và Ialy MR năm 2026 theo Lịch diễn tập số 1067/VHIALY, ngày 07/07/2026","Không phát sinh kiến nghị sau kiểm tra",""] ,
        ["1.4","Thống kê tổng hợp vụ cận nguy xảy ra tại đơn vị","Không phát sinh vụ cận nguy trong kỳ kiểm tra","Không phát sinh kiến nghị sau kiểm tra",""]]},
    {"stt":"2","title":"Hồ sơ, sổ sách, nội quy, quy trình và biện pháp an toàn, sổ ghi biên bản kiểm tra, sổ ghi kiến nghị.","rows":[
        ["2.1","Sổ theo dõi trang cấp BHLĐ; Sổ theo dõi trang bị, dụng cụ an toàn.","Hồ sơ theo dõi trang cấp BHLĐ và dụng cụ an toàn được cập nhật đầy đủ","Không phát sinh kiến nghị sau kiểm tra",""] ,
        ["2.2","Các Quy trình, quy định đã ban hành (liên quan đến công tác an toàn)","Trong tháng không ban hành mới hoặc sửa đổi các quy trình, quy định liên quan đến công tác an toàn","Không phát sinh kiến nghị sau kiểm tra",""]]},
    {"stt":"3","title":"Việc thực hiện các tiêu chuẩn, quy chuẩn, biện pháp an toàn đã ban hành","rows":[
        ["3.1","Thực hiện thao tác theo PTT: Số lượng, kết quả","Các thao tác theo PTT được thực hiện 41 phiếu đúng quy trình, bảo đảm an toàn.","Không phát sinh kiến nghị sau kiểm tra",""] ,
        ["3.2","Tuân thủ, thực hiện đầy đủ thủ tục PCT, LCT: Số lượng, kết quả","Việc thực hiện Phiếu công tác (PCT): 99 PCT, Lệnh công tác (LCT): 49 LCT tại Nhà máy Thủy điện Ialy và Ialy MR: Tỷ lệ PCT không phù hợp là 0%; tỷ lệ LCT không phù hợp là 6,12%.\nQua hậu kiểm phát hiện 03 Lệnh công tác còn tồn tại: Người CHTT không nhập thời gian kết tại cột Kết thúc công tác; Người cấp lệnh không tạo mục “Kiểm tra các biện pháp an toàn trước khi thực hiện lệnh công tác”; Nhân viên đội ĐCT không ký ra khỏi vị trí làm việc.","Đề nghị các cá nhân liên quan tiếp tục rà soát, nâng cao chất lượng lập, kiểm tra Phiếu công tác/Lệnh công tác trước khi phát hành nhằm hạn chế sai sót.",""] ,
        ["3.3","Các biện pháp an toàn","Các biện pháp an toàn được triển khai đầy đủ và hiệu quả","Không phát sinh kiến nghị sau kiểm tra",""]]},
]

# Các mục 4-16 lấy đúng nội dung từ mẫu; có thể chỉnh trực tiếp trên giao diện.
DEFAULT_GROUPS += [
 {"stt":"4","title":"Tình trạng an toàn, vệ sinh của các máy, thiết bị, nhà xưởng, kho tàng và nơi làm việc như: Che chắn tại các vị trí nguy hiểm, độ tin cậy của các cơ cấu an toàn, chống nóng, chống bụi, chiếu sáng, thông gió, thoát nước và các hệ thống khác","rows":[["4.1","Vệ sinh của các máy, thiết bị, nhà xưởng, kho tàng và nơi làm việc như: Che chắn tại các vị trí nguy hiểm, độ tin cậy của các cơ cấu an toàn, chống nóng, chống bụi, chiếu sáng, thông gió, thoát nước và các hệ thống khác","Các khu vực sản xuất được duy trì sạch sẽ; máy móc, thiết bị và hệ thống phụ trợ bảo đảm điều kiện an toàn phục vụ sản xuất.","Không phát sinh kiến nghị sau kiểm tra",""] , ["4.2","Trang bị, phương tiện phục vụ công tác bảo vệ môi trường","Trang bị và phương tiện bảo vệ môi trường đầy đủ, hoạt động tốt","Không phát sinh kiến nghị sau kiểm tra",""] , ["4.3","Việc thu gom, phân loại, xử lý chất thải","Chất thải được thu gom, phân loại và xử lý đúng quy định","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"5","title":"Việc sử dụng, bảo quản trang bị phương tiện bảo vệ cá nhân, phương tiện kỹ thuật phòng cháy chữa cháy, phương tiện cấp cứu y tế","rows":[["5.1","Trang bị phương tiện bảo vệ cá nhân: Sổ theo dõi trang cấp, giao nhận PTBVCN","Trang bị phương tiện bảo vệ cá nhân đầy đủ, hồ sơ theo dõi được cập nhật","Không phát sinh kiến nghị sau kiểm tra",""] , ["5.2","Kiểm tra trang bị phương tiện kỹ thuật phòng cháy chữa cháy","Phương tiện kỹ thuật PCCC được trang bị đầy đủ, tình trạng tốt","Không phát sinh kiến nghị sau kiểm tra",""] , ["5.3","Theo dõi kiểm tra, thử nghiệm, kiểm định các dụng cụ an toàn","Dụng cụ an toàn được kiểm tra, thử nghiệm, kiểm định đúng thời hạn","Không phát sinh kiến nghị sau kiểm tra",""] , ["5.4","Phương tiện cấp cứu y tế","Phương tiện cấp cứu y tế đầy đủ, sẵn sàng sử dụng","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"6","title":"Việc thực hiện các nội dung của kế hoạch ATVSLĐ","rows":[["6.1","Việc thực hiện các nội dung của kế hoạch ATVSLĐ","Các nội dung kế hoạch ATVSLĐ được thực hiện theo tiến độ","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"7","title":"Thực hiện kiến nghị của các đoàn kiểm tra tháng trước","rows":[["7.1","Thực hiện kiến nghị của các đoàn kiểm tra tháng trước","Các kiến nghị của kỳ kiểm tra trước đã được các bộ phận liên quan thực hiện và khắc phục đầy đủ","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"8","title":"Quản lý các thiết bị, vật tư và các chất có yêu cầu nghiêm ngặt về an toàn lao động và các yếu tố nguy hiểm có hại","rows":[["8.1","Quản lý các thiết bị, vật tư và các chất có yêu cầu nghiêm ngặt về an toàn lao động","Thiết bị, vật tư có yêu cầu nghiêm ngặt về ATLĐ được quản lý đúng quy định","Không phát sinh kiến nghị sau kiểm tra",""] , ["8.2","Quản lý các yếu tố nguy hiểm có hại, các khu vực, vị trí làm việc có kết quả quan trắc môi trường không đạt.","Các yếu tố nguy hiểm, có hại được kiểm soát; Qua kiểm tra không phát hiện yếu tố nguy hiểm, có hại vượt mức cho phép hoặc bất thường ảnh hưởng đến an toàn lao động.","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"9","title":"Kiến thức ATVSLĐ, khả năng xử lý sự cố và sơ cứu, cấp cứu của NLĐ","rows":[["9.1","Kiến thức ATVSLĐ, khả năng xử lý sự cố và sơ cứu, cấp cứu của NLĐ","Người lao động nắm vững kiến thức ATVSLĐ và kỹ năng xử lý sự cố","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"10","title":"Tổ chức ăn uống bồi dưỡng, chăm sóc sức khỏe NLĐ","rows":[["10.1","Tổ chức ăn uống bồi dưỡng, chăm sóc sức khỏe NLĐ","Công tác chăm sóc sức khỏe và bồi dưỡng NLĐ được thực hiện đầy đủ","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"11","title":"Hoạt động tự kiểm tra các Kíp, việc khắc phục và giải quyết các đề xuất, kiến nghị về ATVSLĐ","rows":[["11.1","Hoạt động tự kiểm tra các Kíp, việc khắc phục và giải quyết các đề xuất, kiến nghị về ATVSLĐ","Các kíp trực thực hiện đầy đủ việc tự kiểm tra đầu ca, trong ca; Các tồn tại được xử lý kịp thời hoặc báo cáo cấp có thẩm quyền để theo dõi, xử lý theo quy định.","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"12","title":"Trách nhiệm quản lý công tác ATVSLĐ và phong trào quần chúng về ATVSLĐ","rows":[["12.1","Trách nhiệm quản lý công tác ATVSLĐ và phong trào quần chúng về ATVSLĐ","Thực hiện tốt trách nhiệm quản lý ATVSLĐ và phong trào quần chúng","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"13","title":"Công tác sơ kết, tổng kết, báo cáo theo các quy định hiện hành","rows":[["13.1","Công tác sơ kết, tổng kết, báo cáo theo các quy định hiện hành","Trong tháng không phát sinh yêu cầu sơ kết, tổng kết hoặc báo cáo chuyên đề riêng về công tác ATVSLĐ","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"14","title":"Kiểm tra việc xây dựng và thực hiện các phương án ứng cứu khẩn cấp","rows":[["14.1","Kiểm tra việc xây dựng và thực hiện các phương án ứng cứu khẩn cấp","Các phương án ứng cứu khẩn cấp được xây dựng và duy trì thực hiện","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"15","title":"Kiểm tra hiện trường đường dây trên không, trạm điện, nhà máy điện","rows":[["15.1","Kiểm tra hiện trường đường dây trên không, trạm điện, nhà máy điện","Hiện trường đường dây trên không, nhà máy, trạm điện bảo đảm yêu cầu an toàn","Không phát sinh kiến nghị sau kiểm tra",""]]},
 {"stt":"16","title":"Các nội dung khác","rows":[["16.1","Các nội dung khác","Không có nội dung bất thường khác được ghi nhận","Không phát sinh kiến nghị sau kiểm tra",""]]}
]

DEFAULT_MEMBERS = [
 ["Nguyễn Hoàng Phi","Phó Quản đốc - Trưởng đoàn"],
 ["Trần Thanh Chương","QLKT - Người làm Công tác an toàn PX"],
 ["Nguyễn Văn Toàn","Trực ĐKTT"],
 ["Võ Quang Minh","Trực chính TBA 500kV"],
 ["A Ran","Trực trạm 500kV Ialy"],
 ["Trần Thái Hoàng Vũ","Trực phụ máy Gian máy"],
 ["Nguyễn Hồng Quang","Trực chính Gian máy"],
 ["Phùng Ngọc Tú","Trực chính Gian máy"],
]
DEFAULT_RECOMMENDATIONS = [
 "Đề nghị các cá nhân khi thực hiện cấp Lệnh công tác, người cho phép, người chỉ huy trực tiếp và thành viên Đội công tác tăng cường kiểm tra, thực hiện đầy đủ các nội dung trong Lệnh công tác trước khi ký cho phép vào làm việc, trong quá trình thực hiện và khi kết thúc công việc; bảo đảm ghi đầy đủ trình tự công việc, điều kiện an toàn và khóa Lệnh công tác theo đúng quy định.",
 "Các Trưởng ca tăng cường công tác tự kiểm tra đầu ca, trong ca và cuối ca; kịp thời phát hiện, xử lý hoặc báo cáo các nguy cơ mất an toàn, các bất thường của thiết bị và điều kiện làm việc.",
 "Nhân viên vận hành khi thực hiện công việc phải chủ động nhận diện mối nguy, đánh giá rủi ro; trường hợp điều kiện an toàn không bảo đảm phải báo cáo cấp có thẩm quyền để xử lý trước khi tiếp tục công việc.",
 "Người làm công tác an toàn Phân xưởng phối hợp với các Trưởng ca tiếp tục rà soát hồ sơ, biểu mẫu và minh chứng liên quan đến công tác ATVSLĐ; bảo đảm hồ sơ đầy đủ, thống nhất, đúng quy định.",
 "Tiếp tục phổ biến, rút kinh nghiệm các vụ tai nạn lao động, sự cố và vụ cận nguy; nâng cao ý thức chấp hành quy trình, quy định và Văn hóa an toàn của toàn thể nhân viên vận hành./.",
]

def deep_default():
    return {"thang_nam":"07/2026","ngay":"31/07/2026","so_van_ban":"","gio_bd":"08:00","gio_kt":"16:00","ngay_bd":"29/07/2026","ngay_kt":"31/07/2026","members":deepcopy(DEFAULT_MEMBERS),"groups":deepcopy(DEFAULT_GROUPS),"recommendations":deepcopy(DEFAULT_RECOMMENDATIONS),"images":[]}

def save_json(report_id, data):
    (REPORTS_DIR / report_id).mkdir(parents=True, exist_ok=True)
    with open(REPORTS_DIR / report_id / "data.json", "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def load_json(report_id):
    p = REPORTS_DIR / report_id / "data.json"
    if not p.exists(): return None
    with open(p, encoding="utf-8") as f: return json.load(f)

def allowed_file(name):
    return "." in name and name.rsplit(".",1)[1].lower() in ALLOWED_EXT
def prepare_word_image(path):
    """
    Chuẩn bị ảnh để python-docx chèn vào Word.

    JPG/JPEG/PNG/WebP/GIF:
        dùng trực tiếp.

    HEIC/HEIF:
        chuyển sang JPG rồi trả về đường dẫn JPG.
    """

    path = Path(path)

    if path.suffix.lower() not in {".heic", ".heif"}:
        return path

    if not HEIC_SUPPORTED:
        raise RuntimeError(
            "Chưa cài thư viện pillow-heif nên không thể chuyển ảnh HEIC sang JPG."
        )

    jpg_path = path.with_suffix(".jpg")

    with Image.open(path) as img:
        if img.mode != "RGB":
            img = img.convert("RGB")

        img.save(
            jpg_path,
            "JPEG",
            quality=95
        )

    return jpg_path
def replace_text_paragraph(paragraph, text):
    text = str(text or "")
    old = paragraph.runs[0] if paragraph.runs else None
    bold = old.bold if old else None
    italic = old.italic if old else None
    size = old.font.size if old else None
    for r in list(paragraph.runs): paragraph._p.remove(r._r)
    run = paragraph.add_run()
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if size is not None: run.font.size = size
    parts = text.split("\n")
    for i, part in enumerate(parts):
        if i: run.add_break()
        run.add_text(part)

def replace_everywhere(doc, replacements):
    parts = list(doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells: parts += cell.paragraphs
    for s in doc.sections:
        parts += list(s.header.paragraphs) + list(s.footer.paragraphs)
    for p in parts:
        txt = p.text
        new = txt
        for a,b in replacements.items(): new = new.replace(a,b)
        if new != txt: replace_text_paragraph(p,new)

def clear_cell(cell):
    for p in list(cell.paragraphs):
        if p is cell.paragraphs[0]:
            for r in list(p.runs): p._p.remove(r._r)
        else:
            p._element.getparent().remove(p._element)
    if not cell.paragraphs: cell.add_paragraph()

def set_cell(cell, text, bold=False, align=None):
    clear_cell(cell)
    p = cell.paragraphs[0]
    if align is not None: p.alignment = align
    run = p.add_run(str(text or ""))
    run.bold = bold
    for line in str(text or "").split("\n")[1:]:
        run.add_break(); run.add_text(line)

def find_inspection_table(doc):
    for t in doc.tables:
        txt = " ".join(c.text for r in t.rows[:2] for c in r.cells)
        if "Nội dung kiểm tra" in txt and "Kết quả kiểm tra" in txt: return t
    return None

def group_rows(table):
    # Bảng mẫu có 35 dòng. Mỗi dòng cấp 1 có STT nguyên (1..16).
    tops=[]
    for r in table.rows[1:]:
        if re.fullmatch(r"\s*\d+\s*", r.cells[0].text): tops.append(r)
    blocks=[]
    all_rows=list(table.rows)
    for i,top in enumerate(tops):
        nxt=tops[i+1] if i+1<len(tops) else None
        data=[]; started=False
        for r in all_rows:
            if r._tr is top._tr: started=True; continue
            if not started: continue
            if nxt is not None and r._tr is nxt._tr: break
            if re.fullmatch(r"\s*\d+\.\d+\s*",r.cells[0].text): data.append(r)
        # Với mục 6,7,9-16, dòng cấp 1 cũng là dòng dữ liệu.
        merged = bool(top.cells[1]._tc.tcPr.xpath("./w:gridSpan"))
        if not merged:
            data=[top]
        blocks.append((top,data,merged))
    return blocks

def clone_row(table, source, after):
    new=deepcopy(source._tr); after._tr.addnext(new)
    for r in table.rows:
        if r._tr is new: return r
    return table.rows[-1]

def write_inspection_table(doc, groups):
    t=find_inspection_table(doc)
    if not t: return
    blocks=group_rows(t)
    for gi,g in enumerate(groups[:len(blocks)]):
        top,data,merged=blocks[gi]
        if merged:
            set_cell(top.cells[0],g["stt"],bold=True,align=1)
            set_cell(top.cells[1],g["title"],bold=True)
            current=list(data)
        else:
            current=list(data)
        wanted=g.get("rows",[])
        # Nếu mục không có dòng con, tạo đúng một dòng từ chính row hiện tại.
        if not wanted: wanted=[[g["stt"]+".1","","","",""]]
        while len(current)>len(wanted):
            rr=current.pop(); rr._tr.getparent().remove(rr._tr)
        template=current[-1] if current else top
        while len(current)<len(wanted):
            after=current[-1] if current else top
            nr=clone_row(t,template,after); current.append(nr)
        for r, vals in zip(current, wanted):
            for j in range(5):
                set_cell(
                    r.cells[j],
                    vals[j] if j < len(vals) else "",
                    align=1 if j == 0 else None
                )


def find_members_section(doc):
    """
    Tìm phần A. THÀNH PHẦN ĐOÀN KIỂM TRA
    trong nội dung Word.
    """
    paragraphs = doc.paragraphs

    start = None

    for i, p in enumerate(paragraphs):
        text = " ".join(p.text.split()).upper()

        if "THÀNH PHẦN ĐOÀN KIỂM TRA" in text:
            start = i
            break

    if start is None:
        return None

    return start


def write_members(doc, members):
    """
    Cập nhật mục A. THÀNH PHẦN ĐOÀN KIỂM TRA.

    Định dạng:
        [TAB]1. Ông Nguyễn Hoàng Phi       Phó Quản đốc - Trưởng đoàn
        [TAB]2. Ông Trần Thanh Chương      QLKT - Người làm Công tác an toàn PX
        [TAB]3. Ông Nguyễn Văn Toàn        Trực ĐKTT

    - Cột tên thụt vào 1 TAB.
    - Tên/họ tên cỡ chữ 13.
    - Chức vụ cỡ chữ 13.
    - Chức vụ bắt đầu tại một vị trí TAB cố định,
      nên các chức vụ thẳng hàng.
    """

    from docx.oxml import OxmlElement
    from docx.shared import Pt, Cm

    paragraphs = doc.paragraphs
    start = find_members_section(doc)

    if start is None:
        return

    members = members or []

    # =========================================================
    # 1. TÌM CÁC DÒNG THÀNH VIÊN CŨ
    # =========================================================

    member_paragraphs = []

    for i in range(start + 1, len(paragraphs)):
        p = paragraphs[i]

        text = " ".join(p.text.split())

        # Gặp mục tiếp theo thì dừng
        if re.match(r"^[A-Z]\.", text):
            break

        # Nhận diện dòng thành viên cũ
        if re.match(
            r"^\s*\d+\s*[\.\)]\s*Ông\b",
            text,
            re.IGNORECASE
        ):
            member_paragraphs.append(p)

    # =========================================================
    # 2. XÓA TOÀN BỘ DÒNG THÀNH VIÊN CŨ
    # =========================================================

    for p in member_paragraphs:
        p._element.getparent().remove(p._element)

    # =========================================================
    # 3. VỊ TRÍ CHÈN
    # =========================================================

    insert_after = paragraphs[start]._element

    # =========================================================
    # 4. CẤU HÌNH TAB
    # =========================================================

    # Tab thứ nhất:
    # → thụt toàn bộ dòng tên vào 1 TAB
    FIRST_TAB = Cm(1.0)

    # Tab thứ hai:
    # → vị trí bắt đầu của CỘT CHỨC VỤ
    #
    # Anh có thể chỉnh số này:
    # 8.0 cm  : chức vụ gần hơn
    # 9.0 cm  : hiện tại
    # 10.0 cm : chức vụ xa hơn
    #
    ROLE_TAB = Cm(9.0)

    # =========================================================
    # 5. TẠO LẠI DANH SÁCH THÀNH VIÊN
    # =========================================================

    for i, member in enumerate(members):

        name = str(
            member[0] if len(member) > 0 else ""
        ).strip()

        role = str(
            member[1] if len(member) > 1 else ""
        ).strip()

        if not name and not role:
            continue

        # -----------------------------------------------------
        # TẠO PARAGRAPH
        # -----------------------------------------------------

        p = OxmlElement("w:p")

        insert_after.addnext(p)
        insert_after = p

        # -----------------------------------------------------
        # ĐỊNH DẠNG PARAGRAPH
        # -----------------------------------------------------

        pPr = OxmlElement("w:pPr")

        # Tab stops
        tabs = OxmlElement("w:tabs")

        # TAB 1 - thụt tên vào 1 tab
        tab1 = OxmlElement("w:tab")
        tab1.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
            "left"
        )
        tab1.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pos",
            str(int(FIRST_TAB.cm * 567))
        )
        tabs.append(tab1)

        # TAB 2 - cột chức vụ cố định
        tab2 = OxmlElement("w:tab")
        tab2.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
            "left"
        )
        tab2.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pos",
            str(int(ROLE_TAB.cm * 567))
        )
        tabs.append(tab2)

        pPr.append(tabs)

        # Khoảng cách dòng
        spacing = OxmlElement("w:spacing")
        spacing.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}after",
            "0"
        )
        spacing.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}line",
            "240"
        )

        pPr.append(spacing)

        p.insert(0, pPr)

        # =====================================================
        # TAB 1 - THỤT VÀO 1 TAB
        # =====================================================

        rtab1 = OxmlElement("w:r")
        tab_run1 = OxmlElement("w:tab")
        rtab1.append(tab_run1)
        p.append(rtab1)

        # =====================================================
        # SỐ + HỌ TÊN
        # =====================================================

        r1 = OxmlElement("w:r")

        rpr1 = OxmlElement("w:rPr")

        # Font Times New Roman
        rfonts1 = OxmlElement("w:rFonts")

        rfonts1.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
            "Times New Roman"
        )

        rfonts1.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
            "Times New Roman"
        )

        rpr1.append(rfonts1)

        # Cỡ chữ 13
        size1 = OxmlElement("w:sz")
        size1.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
            "26"
        )
        rpr1.append(size1)

        sizecs1 = OxmlElement("w:szCs")
        sizecs1.set(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
            "26"
        )
        rpr1.append(sizecs1)

        r1.append(rpr1)

        t1 = OxmlElement("w:t")

        t1.text = f"{i + 1}. Ông {name}"

        r1.append(t1)
        p.append(r1)

        # =====================================================
        # TAB ĐẾN CỘT CHỨC VỤ
        # =====================================================

        if role:

            rtab2 = OxmlElement("w:r")

            tab_run2 = OxmlElement("w:tab")

            rtab2.append(tab_run2)

            p.append(rtab2)

            # =================================================
            # CHỨC VỤ
            # =================================================

            r2 = OxmlElement("w:r")

            rpr2 = OxmlElement("w:rPr")

            rfonts2 = OxmlElement("w:rFonts")

            rfonts2.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii",
                "Times New Roman"
            )

            rfonts2.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi",
                "Times New Roman"
            )

            rpr2.append(rfonts2)

            # Cỡ chữ 13
            size2 = OxmlElement("w:sz")
            size2.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                "26"
            )

            rpr2.append(size2)

            sizecs2 = OxmlElement("w:szCs")
            sizecs2.set(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val",
                "26"
            )

            rpr2.append(sizecs2)

            r2.append(rpr2)

            t2 = OxmlElement("w:t")

            t2.text = role

            r2.append(t2)

            p.append(r2)
def find_signature_table(doc):
    for t in doc.tables:
        txt=" ".join(c.text for r in t.rows[:1] for c in r.cells)
        if "THÀNH VIÊN ĐOÀN KIỂM TRA" in txt and "TRƯỞNG ĐOÀN" in txt: return t
    return None

def write_signatures(doc, members):
    """
    Phần ký:
    - Bỏ chức vụ, chỉ lấy họ tên.
    - Thành viên chia thành 2 cột.
    - Người 1 ở cột trái, người 2 ở cột phải cùng một hàng.
    - Người 3 ở cột trái, người 4 ở cột phải cùng một hàng...
    - Toàn bộ cụm thành viên nằm sát về bên trái.
    - Trưởng đoàn giữ ở cột bên phải.
    """

    t = find_signature_table(doc)

    if not t or len(t.columns) < 3:
        return

    members = members or []

    # =========================================================
    # 1. TÌM TRƯỞNG ĐOÀN
    # =========================================================

    leader_index = None

    for i, member in enumerate(members):
        role = str(member[1] if len(member) > 1 else "").strip()

        if "trưởng đoàn" in role.lower():
            leader_index = i
            break

    if leader_index is None and members:
        leader_index = 0

    leader = members[leader_index] if leader_index is not None else ["", ""]

    # Danh sách thành viên, bỏ trưởng đoàn
    member_list = [
        member
        for i, member in enumerate(members)
        if i != leader_index
    ]

    # =========================================================
    # 2. XÓA NỘI DUNG CŨ
    # =========================================================

    while len(t.rows) < 2:
        t.add_row()

    for row in t.rows[1:]:
        for cell in row.cells:
            clear_cell(cell)

    # =========================================================
    # 3. CỘT THÀNH VIÊN BÊN TRÁI
    # =========================================================

    left_cells = [
        row.cells[1]
        for row in t.rows[1:]
    ]

    if left_cells:

        left = left_cells[0]

        # Gộp toàn bộ vùng thành viên thành một ô
        for cell in left_cells[1:]:
            left = left.merge(cell)

        clear_cell(left)

        from docx.shared import Cm, Pt

        # -----------------------------------------------------
        # TẠO 2 CỘT BẰNG TAB
        # -----------------------------------------------------

        # Khoảng cách giữa cột 1 và cột 2.
        # ANH CHỈNH SỐ Ở ĐÂY NẾU MUỐN:
        # 6.0 = cột 2 gần cột 1 hơn
        # 7.0 = hiện tại
        # 8.0 = cột 2 xa hơn
        COLUMN_2_POSITION = Cm(7.0)

        # Chia thành từng hàng, mỗi hàng 2 người
        for row_index in range(0, len(member_list), 2):

            p = (
                left.paragraphs[0]
                if row_index == 0
                and len(left.paragraphs) == 1
                and not left.paragraphs[0].text
                else left.add_paragraph()
            )

            p.alignment = 0

            # =================================================
            # NGƯỜI THỨ NHẤT - CỘT TRÁI
            # =================================================

            member1 = member_list[row_index]

            name1 = str(
                member1[0] if len(member1) > 0 else ""
            ).strip()

            if name1:

                run = p.add_run(
                    f"{row_index + 1}. Ông: {name1}"
                )

                run.bold = False
                run.font.size = Pt(11)

            # =================================================
            # NGƯỜI THỨ HAI - CỘT PHẢI
            # =================================================

            if row_index + 1 < len(member_list):

                member2 = member_list[row_index + 1]

                name2 = str(
                    member2[0] if len(member2) > 0 else ""
                ).strip()

                if name2:

                    # Tạo tab đến vị trí cột 2
                    p.paragraph_format.tab_stops.add_tab_stop(
                        COLUMN_2_POSITION
                    )

                    p.add_run("\t")

                    run = p.add_run(
                        f"{row_index + 2}. Ông: {name2}"
                    )

                    run.bold = False
                    run.font.size = Pt(11)

            # =================================================
            # KHOẢNG TRỐNG ĐỂ KÝ
            # =================================================

            for _ in range(4):
                p.add_run().add_break()

            p.paragraph_format.space_after = Pt(2)

    # =========================================================
    # 4. TRƯỞNG ĐOÀN - CỘT BÊN PHẢI
    # =========================================================

    right_cells = [
        row.cells[2]
        for row in t.rows[1:]
    ]

    if right_cells:

        right = right_cells[0]

        for cell in right_cells[1:]:
            right = right.merge(cell)

        clear_cell(right)

        name = str(
            leader[0] if len(leader) > 0 else ""
        ).strip()

        if name:

            from docx.shared import Pt

            p = right.paragraphs[0]
            p.alignment = 1

            # Khoảng trống để ký
            for _ in range(5):
                p.add_run().add_break()

            # Chỉ ghi tên, không ghi chức vụ
            run = p.add_run(
                "Ông: " + name
            )

            run.bold = True
            run.font.size = Pt(11)

    # =========================================================
    # 5. XÓA CỘT NƠI NHẬN
    # =========================================================

    for row in t.rows[2:]:
        clear_cell(row.cells[0])
def find_appendix_table(doc):
    for t in doc.tables:
        txt=" ".join(c.text for r in t.rows[:2] for c in r.cells)
        if "Hình ảnh hiện trường NMTĐ Ialy" in txt and "Hình ảnh hiện trường NMTĐ Ialy MR" in txt: return t
    return None

def clear_drawing(cell):
    for d in list(cell._tc.xpath(".//w:drawing")): d.getparent().remove(d)

def write_images(doc, images):
    t = find_appendix_table(doc)

    if not t:
        return

    # Mỗi STT dùng 2 dòng:
    # dòng 1: STT + tiêu đề ảnh
    # dòng 2: ảnh
    needed = max(1, len(images), 6)

    # Đảm bảo đủ số dòng cho số lượng ảnh cần chèn
    while len(t.rows) < 1 + needed * 2:
        cap = deepcopy(t.rows[-2]._tr)
        img = deepcopy(t.rows[-1]._tr)

        t.rows[-1]._tr.addnext(cap)
        t.rows[-1]._tr.addnext(img)

    # Xóa nội dung ảnh/caption cũ
    for i in range(needed):
        cr = 1 + i * 2
        ir = cr + 1

        for col in (0, 1, 2, 3):
            clear_cell(t.rows[cr].cells[col])
            clear_drawing(t.rows[ir].cells[col])

    # Ghi ảnh mới
    for item in images:

        try:
            stt = int(item.get("stt", 1))
        except Exception:
            stt = 1

        side = item.get("side", "ST")

        # Cột:
        # ST -> cột 1
        # MR -> cột 3
        col = 1 if side == "ST" else 3

        cr = 1 + (stt - 1) * 2
        ir = cr + 1

        if cr >= len(t.rows):
            continue

        # -------------------------------
        # CAPTION
        # -------------------------------
        set_cell(
            t.rows[cr].cells[col],
            item.get("caption", ""),
            align=1
        )

        # -------------------------------
        # STT
        # -------------------------------
        set_cell(
            t.rows[cr].cells[col - 1],
            str(stt),
            bold=True,
            align=1
        )

        # -------------------------------
        # ẢNH
        # -------------------------------
        path = item.get("path")

        if not path:
            continue

        path = Path(path)

        if not path.exists():
            continue

        try:
            # HEIC/HEIF -> JPG
            word_image = prepare_word_image(path)

            p = t.rows[ir].cells[col].paragraphs[0]
            p.alignment = 1

            run = p.add_run()

            run.add_picture(
                str(word_image),
                width=Inches(3.8)
            )

        except Exception as e:
            print(
                f"[WARNING] Không chèn được ảnh "
                f"{path}: {type(e).__name__}: {e}"
            )

def build_doc(data, report_id):
    if not TEMPLATE.exists(): raise FileNotFoundError(f"Thiếu file mẫu: {TEMPLATE.name}")
    doc=Document(str(TEMPLATE))
    thang=data.get("thang_nam","07/2026")
    ngay=data.get("ngay","31/07/2026")
    replacements={
      "THÁNG 7/2026":f"THÁNG {thang}",
      "tháng 7/2026":f"tháng {thang}",
      "tháng 7 năm 2026":f"tháng {thang.split('/')[0]} năm {thang.split('/')[-1]}",
      "ngày 31 tháng 07 năm 2026":f"ngày {ngay.replace('/', ' tháng ')}",
      "Số: /VHIALY":f"Số: {data.get('so_van_ban','')} /VHIALY" if data.get('so_van_ban','') else "Số: /VHIALY",
      "Từ 08 giờ 00 phút đến 16 giờ 00 phút, trong các ngày từ 29/07/2026 đến 31/07/2026.":f"Từ {data.get('gio_bd','08:00')} đến {data.get('gio_kt','16:00')}, trong các ngày từ {data.get('ngay_bd','29/07/2026')} đến {data.get('ngay_kt','31/07/2026')}."
    }
    replace_everywhere(doc,replacements)

# Cập nhật danh sách thành phần đoàn kiểm tra ở mục A
    write_members(doc,data.get("members",[]))

# Cập nhật bảng nội dung kiểm tra
    write_inspection_table(doc,data.get("groups",[]))

# Cập nhật phần ký
    write_signatures(doc,data.get("members",[]))
    images=[]
    base=REPORTS_DIR/report_id
    for x in data.get("images",[]):
        y=dict(x); y["path"]=str(base/x["filename"]); images.append(y)
    write_images(doc,images)
    out=OUTPUT_DIR/f"Bien_ban_ATVSLD_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(out)); return out


def parse_form(form):
    data=deep_default()
    for key in ["thang_nam","ngay","so_van_ban","gio_bd","gio_kt","ngay_bd","ngay_kt"]:
        data[key]=form.get(key,"").strip()
    names=form.getlist("member_name"); roles=form.getlist("member_role")
    data["members"]=[[names[i].strip() if i<len(names) else "",roles[i].strip() if i<len(roles) else ""] for i in range(max(len(names),len(roles))) if (names[i].strip() if i<len(names) else "") or (roles[i].strip() if i<len(roles) else "")]
    groups=[]
    stts=form.getlist("group_stt"); titles=form.getlist("group_title")
    for gi,stt in enumerate(stts):
        pre=f"g{gi}_"; idx=form.getlist(pre+"idx"); c=form.getlist(pre+"content"); r=form.getlist(pre+"result"); k=form.getlist(pre+"recommendation"); n=form.getlist(pre+"note")
        rows=[]
        for i in range(len(idx)):
            rows.append([idx[i],c[i] if i<len(c) else "",r[i] if i<len(r) else "",k[i] if i<len(k) else "",n[i] if i<len(n) else ""])
        groups.append({"stt":stt,"title":titles[gi] if gi<len(titles) else "","rows":rows})
    data["groups"]=groups
    data["recommendations"]=[x.strip() for x in form.getlist("recommendation") if x.strip()]
    data["images"]=[]
    return data


@app.route("/")
def index():
    return redirect(url_for("create_report"))

@app.route("/tao-bao-cao",methods=["GET","POST"])
@login_required
def create_report():
    if not current_user.can_edit:
        flash("Anh/chị không có quyền tạo báo cáo.")
        return redirect(url_for("history"))
    if request.method=="GET":
        return render_template("index.html",data=deep_default(),edit_id="")
    data=parse_form(request.form)
    rid=datetime.now().strftime("%Y%m%d_%H%M%S_")+uuid.uuid4().hex[:6]
    folder=REPORTS_DIR/rid; folder.mkdir(parents=True,exist_ok=True)
    # Nhận file theo từng ô: image_stt, image_side, image_caption.
    files=request.files.getlist("image_file"); stts=request.form.getlist("image_stt"); sides=request.form.getlist("image_side"); caps=request.form.getlist("image_caption")
    for i,f in enumerate(files):
        if not f or not f.filename or not allowed_file(f.filename): continue
        fn=f"{uuid.uuid4().hex}_{Path(f.filename).name}"
        f.save(folder/fn)
        data["images"].append({"filename":fn,"stt":int(stts[i]) if i<len(stts) and stts[i].isdigit() else 1,"side":sides[i] if i<len(sides) else "ST","caption":caps[i] if i<len(caps) else ""})
    data["created_at"]=datetime.now().strftime("%d/%m/%Y %H:%M")
    data["updated_at"]=data["created_at"]
    save_json(rid,data)
    try: out=build_doc(data,rid)
    except Exception as e: flash(f"Lỗi xuất Word: {type(e).__name__}: {e}"); return render_template("index.html",data=data,edit_id=""),500
    return send_file(out,as_attachment=True,download_name=out.name)

@app.route("/lich-su")
def history():
    items=[]
    for p in REPORTS_DIR.iterdir():
        if p.is_dir() and (p/"data.json").exists():
            d=load_json(p.name); items.append({"id":p.name,**d})
    items.sort(key=lambda x:x.get("updated_at",x["id"]),reverse=True)
    return render_template("history.html",items=items)

@app.route("/sua/<rid>")
@login_required
def edit_report(rid):
    if not current_user.is_authenticated or not current_user.can_edit:
        flash("Anh/chị không có quyền sửa báo cáo.")
        return redirect(url_for("history"))

    data = load_json(rid)

    if not data:
        flash("Không tìm thấy báo cáo.")
        return redirect(url_for("history"))

    return render_template("index.html", data=data, edit_id=rid)

@app.route("/cap-nhat/<rid>",methods=["POST"])
def update_report(rid):
    old=load_json(rid)
    if not old: flash("Không tìm thấy báo cáo."); return redirect(url_for("history"))
    data=parse_form(request.form)
    data["images"]=old.get("images",[])
    folder=REPORTS_DIR/rid; folder.mkdir(exist_ok=True)
    files=request.files.getlist("image_file"); stts=request.form.getlist("image_stt"); sides=request.form.getlist("image_side"); caps=request.form.getlist("image_caption")
    # Nếu người dùng chọn file mới ở ô nào thì thay ảnh cũ cùng STT + side.
    for i,f in enumerate(files):
        if not f or not f.filename or not allowed_file(f.filename): continue
        stt=int(stts[i]) if i<len(stts) and stts[i].isdigit() else 1; side=sides[i] if i<len(sides) else "ST"
        for oldimg in list(data["images"]):
            if int(oldimg.get("stt",0))==stt and oldimg.get("side")==side:
                try: (folder/oldimg["filename"]).unlink(missing_ok=True)
                except Exception: pass
                data["images"].remove(oldimg)
        fn=f"{uuid.uuid4().hex}_{Path(f.filename).name}"; f.save(folder/fn)
        data["images"].append({"filename":fn,"stt":stt,"side":side,"caption":caps[i] if i<len(caps) else ""})
    data["created_at"]=old.get("created_at",datetime.now().strftime("%d/%m/%Y %H:%M")); data["updated_at"]=datetime.now().strftime("%d/%m/%Y %H:%M")
    save_json(rid,data)
    try: out=build_doc(data,rid)
    except Exception as e: flash(f"Lỗi cập nhật Word: {type(e).__name__}: {e}"); return render_template("index.html",data=data,edit_id=rid),500
    return send_file(out,as_attachment=True,download_name=out.name)

@app.route("/xoa/<rid>", methods=["POST"])
@login_required
def delete_report(rid):
    if not current_user.can_edit:
        flash("Anh/chị không có quyền xóa báo cáo.")
        return redirect(url_for("history"))

    folder = REPORTS_DIR / rid

    if folder.exists():
        shutil.rmtree(folder)

    flash("Đã xóa báo cáo.")
    return redirect(url_for("history"))

@app.route("/word/<rid>")
def open_word(rid):
    data = load_json(rid)

    if not data:
        flash("Không tìm thấy báo cáo.")
        return redirect(url_for("history"))

    try:
        out = build_doc(data, rid)

        return send_file(
            out,
            as_attachment=False,
            download_name=out.name
        )

    except Exception as e:
        flash(
            f"Lỗi mở Word: {type(e).__name__}: {e}"
        )
        return redirect(url_for("history"))
@app.route("/report-image/<rid>/<filename>")
def report_image(rid,filename):
    return send_from_directory(REPORTS_DIR/rid,filename)

ensure_admin()
if __name__=="__main__":
    # 0.0.0.0 để máy khác trong LAN truy cập được.
    app.run(host="0.0.0.0",port=5001,debug=False)
